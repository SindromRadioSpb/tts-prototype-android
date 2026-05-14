# scripts/research/gen_smoke_check_docx.py
# Generates Smoke-check/SMOKE_CHECK_RESEARCH_MODE_v3_2_0.docx
# Run: python scripts/research/gen_smoke_check_docx.py

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


import sys
REPO = Path(__file__).resolve().parents[2]
OUT_DIR = REPO / "Smoke-check"
# Allow override via CLI arg in case Word holds a lock on the canonical file.
_default_out = OUT_DIR / "SMOKE_CHECK_RESEARCH_MODE_v3_2_0.docx"
OUT_FILE = Path(sys.argv[1]) if len(sys.argv) > 1 else _default_out

DOC_TITLE = "Smoke-check: Research Mode (Direction 11B) — v3.2.0"
DOC_SUBTITLE = "Ручная проверка перед deployment в реальной группе ulpan"
DOC_VERSION = "v3.2.0 (tag 32d8cb4) · smoke revision r3 2026-05-14"
DOC_DATE = date.today().isoformat()


# ─────────────────────────────────────────────────────────────────────────────
# Style helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, size=None, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items, bullet_style='List Bullet'):
    for it in items:
        p = doc.add_paragraph(style=bullet_style)
        run = p.add_run(it)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(it)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    # light-grey shading via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ─────────────────────────────────────────────────────────────────────────────
# Test-case rendering
# ─────────────────────────────────────────────────────────────────────────────

def add_test_case(doc, tc_id, title, *,
                  preconditions=None,
                  steps=None,
                  expected=None,
                  notes=None,
                  refs=None):
    """Render a single test case as a table-styled block."""
    p = doc.add_paragraph()
    run = p.add_run(f"{tc_id} — {title}")
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(13.5)

    def row(label, content_lines):
        r = table.add_row()
        c0 = r.cells[0]
        c1 = r.cells[1]
        c0.width = Cm(3.2)
        c1.width = Cm(13.5)
        set_cell_shading(c0, 'EAF1F8')
        c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = c0.paragraphs[0]
        rr = p0.add_run(label)
        rr.bold = True
        rr.font.size = Pt(10)

        # first paragraph already in cell — re-use for first line, add more if needed
        first = True
        for line in content_lines:
            if first:
                pp = c1.paragraphs[0]
                first = False
            else:
                pp = c1.add_paragraph()
            r1 = pp.add_run(line)
            r1.font.name = 'Calibri'
            r1.font.size = Pt(10)

    if preconditions:
        row("Предусловия", preconditions if isinstance(preconditions, list) else [preconditions])
    if steps:
        # number them
        numbered = [f"{i+1}. {s}" for i, s in enumerate(steps)]
        row("Шаги", numbered)
    if expected:
        row("Ожидаемый результат", expected if isinstance(expected, list) else [expected])
    if refs:
        row("Refs", refs if isinstance(refs, list) else [refs])

    # PASS/FAIL/N-A row
    r = table.add_row()
    c0 = r.cells[0]
    c1 = r.cells[1]
    set_cell_shading(c0, 'EAF1F8')
    p0 = c0.paragraphs[0]
    rr = p0.add_run("Результат")
    rr.bold = True
    rr.font.size = Pt(10)
    pp = c1.paragraphs[0]
    for label in ["☐ PASS", "    ☐ FAIL", "    ☐ N/A"]:
        rs = pp.add_run(label + "        ")
        rs.font.name = 'Calibri'
        rs.font.size = Pt(10.5)
    note_p = c1.add_paragraph()
    note_run = note_p.add_run("Заметки / скриншот: ")
    note_run.font.size = Pt(10)
    note_run.italic = True
    note_p.add_run("_" * 80)

    # spacing after
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Document body
# ─────────────────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Cover ────────────────────────────────────────────────────────────────
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_title.add_run(DOC_TITLE)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_sub.add_run(DOC_SUBTITLE)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_rows = [
        ("Версия приложения", DOC_VERSION),
        ("Дата выпуска smoke-check", DOC_DATE),
        ("Аудитория", "Tester (UI only) + Developer (DevTools/CLI)"),
        ("Длительность прогона", "~2 ч (Tester) + ~1.5 ч (Developer)"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta_table.cell(i, 0).text = k
        meta_table.cell(i, 1).text = v
        set_cell_shading(meta_table.cell(i, 0), 'EAF1F8')
        for p in meta_table.cell(i, 0).paragraphs:
            for rr in p.runs:
                rr.bold = True

    doc.add_paragraph()
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note_p.add_run(
        "Документ — обязательный pre-deployment gate. "
        "Прогон полностью green = разрешение на distribution когорты."
    )
    rn.italic = True
    rn.font.size = Pt(10.5)
    rn.font.color.rgb = RGBColor(0x88, 0x44, 0x00)

    add_page_break(doc)

    # ── Section 0: Pre-flight ────────────────────────────────────────────────
    add_heading(doc, "Раздел 0. Pre-flight (общее)", level=1,
                color=RGBColor(0x1F, 0x4E, 0x79))

    add_heading(doc, "0.1 Инвалидные ситуации (skip smoke если применимо)", level=2)
    add_bullets(doc, [
        "Working tree не на тэге v3.2.0 (commit 32d8cb4). Smoke имеет смысл только на этом snapshot или новее.",
        "RESEARCH_DATA_DIR указывает на production-данные. Smoke должен идти на чистый временный каталог.",
        "На сервере уже есть когорта с production-студентами. Используй staging-deploy или local dev.",
    ])

    add_heading(doc, "0.2 Подготовка стенда", level=2)
    add_numbered(doc, [
        "Подними чистый локальный стенд: `git checkout v3.2.0` (или новее на main).",
        "Удали/переименуй существующий `data/research/`, если есть.",
        "Установи playwright chromium один раз: `npx playwright install chromium`.",
        "Запусти автосмоук как baseline (должно быть 60/60 + 9 PNG):  `npm run smoke:research`. Если первый прогон fail на teacher-smoke (Timeout на .summary-tile) — это известный flake, перепрогон должен быть green. См. §0.6.",
        "Запусти dev-сервер: `npm start`. Открой `http://localhost:3000`. Если порт занят — `set PORT=3001 && npm start` (Windows) или `PORT=3001 npm start` (POSIX).",
        "Открой второй таб: `http://localhost:3000/teacher.html` (пока пусто).",
    ])

    add_heading(doc, "0.3 Provision test artifacts (ОДНОЙ КОМАНДОЙ — ОБЯЗАТЕЛЬНО)", level=2)
    add_para(doc,
        "Все тестовые когорты, fake-данные и CSV-фикстуры готовятся одной командой. "
        "Tester и Developer оба используют одинаковые артефакты — никаких ручных шагов provisioning'а нет.",
        size=10.5)

    add_code(doc, "npm run smoke:prep")

    add_para(doc, "Что создаст команда:", bold=True, size=11)
    add_bullets(doc, [
        "data/research/TEST-PILOT-A/ — пустая когорта для TC-T-5 (consent + join).",
        "data/research/TEST-PILOT-B/ — пустая когорта для TC-T-9 (смена когорты).",
        "data/research/SEED-K5/ — seeded когорта с 12 fake students × 14 days, k≥5 met. Для TC-T-18 (заполненный teacher dashboard).",
        "Smoke-check/outcomes-good.csv — валидный CSV с реальными UUID из SEED-K5 (12 rows). Для TC-T-19.",
        "Smoke-check/outcomes-bad-no-header.csv — sad path: нет header row. Для TC-T-20a.",
        "Smoke-check/outcomes-bad-empty.csv — sad path: header без data rows. Для TC-T-20b.",
        "Smoke-check/test-cohort-credentials.txt — все researcher tokens + 5 первых UUID. Готовый к использованию.",
    ])
    add_para(doc, "Cleanup после прогона smoke:", bold=True, size=11)
    add_code(doc, "npm run smoke:prep:clean")
    add_para(doc,
        "Все cohort-директории + все Smoke-check/outcomes-*.csv + credentials.txt — удалены. Skript идемпотентен; повторный запуск smoke:prep всегда даст свежие токены и UUID.",
        italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55))

    add_heading(doc, "0.4 Инструменты тестера / разработчика", level=2)
    add_bullets(doc, [
        "Tester: только браузер. Chromium / Firefox / Safari mobile.",
        "Developer: DevTools (Network, Application/Storage, Console), curl, текстовый редактор.",
        "Обе версии: ручка/блокнот или digital-таб для заметок и скриншотов.",
        "Tester должен иметь под рукой `Smoke-check/test-cohort-credentials.txt` (для TC-T-5, T-9, T-17).",
    ])

    add_heading(doc, "0.5 Окно прогона", level=2)
    add_para(doc,
        "Запускай smoke на изолированном профиле/инкогнито-окне, чтобы localStorage "
        "приложения не сталкивался с твоим обычным состоянием. После прогона "
        "удали обе тестовые когорты: `rm -rf data/research/TEST-PILOT-A data/research/TEST-PILOT-B`.")

    add_heading(doc, "0.6 Известные особенности (read me)", level=2)
    add_bullets(doc, [
        "Flaky teacher-smoke: первый прогон `npm run smoke:research` после простоя сервера иногда падает с Timeout 10000ms на `.summary-tile`. Перепрогон green. Не считай первый fail регрессией — проверь повторно.",
        "Cohort code limit: maxlength=16 в UI; regex `[A-Z0-9-]{4,16}`. Не используй имена длиннее 16 символов.",
        "Researcher token printed ONCE: в stdout `create_cohort.js`. На диске лежит только sha256 hash. Если потерял plaintext — Procedure B rotation (TC-D-18) даёт новый.",
        "👁 «Что собрано» показывает лог отправленных uploads (до 30 записей), не текущее состояние черновика. До первого upload экран пустой с пометкой «Пока ничего не отправлено».",
        "JSONL deletions audit log является append-only — после TC-D-10/20 строка не очищается, она добавляется.",
    ])

    add_heading(doc, "0.7 Acceptance summary (заполнить в конце)", level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    hdr[0].text = "Раздел"
    hdr[1].text = "Тестов"
    hdr[2].text = "PASS / FAIL / N-A"
    body_rows = [
        ("Часть I — Tester", "21", "____ / ____ / ____"),
        ("Часть II — Developer", "25", "____ / ____ / ____"),
        ("Auto-smoke (npm run smoke:research)", "60 + 9 PNG", "☐ green  ☐ red"),
        ("Итог: разрешение на pilot-deploy", "—", "☐ ДА  ☐ НЕТ"),
    ]
    for row in body_rows:
        r = t.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────
    # Часть I — Tester
    # ─────────────────────────────────────────────────────────────────────
    add_heading(doc, "Часть I — Tester (UI only, без DevTools)",
                level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    add_para(doc,
        "Прогон рассчитан на нетехнического тестера. Действуй строго по шагам. "
        "Если шаг невозможно выполнить — отмечай FAIL и фиксируй детали в Заметках. "
        "Скриншот делай на каждом «Ожидаемый результат» хотя бы одной строкой.",
        italic=True, size=10.5, color=RGBColor(0x55, 0x55, 0x55))

    add_test_case(doc, "TC-T-1", "Открытие приложения и поиск 📊 Research mode",
        preconditions=[
            "Чистый профиль браузера (или режим инкогнито).",
            "Стенд запущен на http://localhost:3000.",
        ],
        steps=[
            "Открой http://localhost:3000.",
            "Дождись полной загрузки (видна верхняя панель с иконками).",
            "Найди иконку 📊 в верхней панели.",
            "Кликни на 📊.",
        ],
        expected=[
            "Открывается модал «📊 Research mode».",
            "Статус: «Выключен» (зелёная отметка отсутствует).",
            "Видна большая кнопка «📊 Дать согласие на участие».",
            "Внизу — пометка про anonymity / privacy.",
        ])

    add_test_case(doc, "TC-T-2", "Просмотр consent screen",
        preconditions="Открыт модал Research mode (TC-T-1).",
        steps=[
            "Кликни «📊 Дать согласие на участие».",
            "Прочитай содержимое окна (что собирается / что НЕ собирается / права).",
            "Внутри найди свёртку «Полный текст согласия».",
        ],
        expected=[
            "Заголовок «📋 Согласие на участие».",
            "Видны 3 секции: «Что мы собираем», «Что НЕ собирается никогда», «Ваши права».",
            "Внизу — список из 5 чекбоксов согласия.",
            "Версия документа: 1.0.",
            "Кнопки: «Принимаю» / «Отмена».",
        ])

    add_test_case(doc, "TC-T-3", "Попытка принять без всех чекбоксов",
        preconditions="Открыт consent screen (TC-T-2).",
        steps=[
            "Отметь только 4 из 5 чекбоксов.",
            "Нажми «Принимаю».",
        ],
        expected=[
            "Появляется toast: «Отметьте все 5 пунктов согласия».",
            "Модал остаётся открытым (не закрылся).",
            "Состояние Research mode не сменилось на «Активен».",
        ])

    add_test_case(doc, "TC-T-4", "Принятие consent",
        preconditions="Открыт consent screen.",
        steps=[
            "Отметь все 5 чекбоксов.",
            "Нажми «Принимаю».",
        ],
        expected=[
            "Появляется toast: «Согласие сохранено».",
            "Модал закрывается; повторно открой 📊.",
            "Статус: «Без cohort code» (оранжевая отметка) ИЛИ «Активен».",
            "В строке Anonymous student ID — UUID вида 8-4-4-4-12 hex.",
            "Кнопка «🔗 Присоединиться к когорте» доступна.",
        ])

    add_test_case(doc, "TC-T-5", "Присоединение к когорте",
        preconditions=[
            "consent принят (TC-T-4).",
            "На сервере провижионирована тестовая когорта TEST-PILOT-A (Pre-flight §0.3).",
        ],
        steps=[
            "Открой 📊 → «🔗 Присоединиться к когорте».",
            "Введи cohort code: TEST-PILOT-A (12 chars, в пределах maxlength=16).",
            "Нажми «Присоединиться».",
        ],
        expected=[
            "Toast: «✓ Вы в когорте: TEST-PILOT-A».",
            "Модал перерисовывается; в панели: Cohort code = TEST-PILOT-A.",
            "Видны 5 кнопок действий: 🔁 Сменить когорту, 👁 Что собрано, ⬆ Отправить сейчас, 🎓 Сдать экзамен, 🗑 Отозвать согласие.",
        ])

    add_test_case(doc, "TC-T-6", "Транспаренси — «👁 Что собрано» (empty state)",
        preconditions="Активная когорта (TC-T-5). До первого upload!",
        steps=[
            "Открой 📊 → «👁 Что собрано».",
        ],
        expected=[
            "Открывается модал «👁 Что собрано».",
            "Сверху — пояснение: «Здесь видны все последние uploads (до 30). Каждая запись — это то, что было отправлено за этот день.»",
            "По центру — empty state: «Пока ничего не отправлено. Uploads появятся после первого полного дня активности.»",
            "Внизу — пометка: «Сохраняются последние 30 uploads. Старые записи отбрасываются.»",
            "Это НЕ показывает текущий черновик агрегации (by design — превью будущего upload отсутствует в v3.2). После TC-T-7 первый upload появится строкой.",
        ],
        refs=["Реализация: public/js/research-ui.js function openTransparency() — uploads-only лог."])

    add_test_case(doc, "TC-T-7", "Ручной upload",
        preconditions="Активная когорта.",
        steps=[
            "Используй приложение 1-2 минуты (открой пару экранов — это породит хотя бы один heartbeat-event).",
            "Открой 📊 → «⬆ Отправить сейчас».",
            "Дождись завершения операции (toast/индикатор).",
            "Открой «👁 Что собрано» повторно.",
        ],
        expected=[
            "Toast после upload: «Отправлено» / похожий success.",
            "В лог-таблице транспаренси появляется первая строка: Дата + Статус «✓ stored» + столбцы Минут/SRS/Заметок/Bytes.",
            "Если activity не было — допустим статус «↻ dedupe» или skip; повторный upload ничего нового не покажет (см. TC-D-7).",
            "Поле «Последний upload» в основной панели обновилось.",
        ])

    add_test_case(doc, "TC-T-8", "Self-report экзамена",
        preconditions="Активная когорта.",
        steps=[
            "Открой 📊 → «🎓 Сдать экзамен».",
            "Введи балл: 78.",
            "Опционально введи confidence: 4.",
            "Нажми «Отправить».",
        ],
        expected=[
            "Toast: «Outcome сохранён».",
            "В лог uploads попадает upload с outcome populated.",
            "Модал закрывается.",
        ])

    add_test_case(doc, "TC-T-9", "Смена когорты",
        preconditions=[
            "Активная когорта TEST-PILOT-A.",
            "В Pre-flight §0.3 провижионирована вторая когорта TEST-PILOT-B.",
        ],
        steps=[
            "Открой 📊 → «🔁 Сменить когорту».",
            "Поле ввода уже заполнено текущим кодом TEST-PILOT-A. Очисти поле.",
            "Введи: TEST-PILOT-B (12 chars).",
            "Нажми «Присоединиться».",
        ],
        expected=[
            "Toast: «✓ Вы в когорте: TEST-PILOT-B».",
            "В панели Cohort code = TEST-PILOT-B.",
            "Anonymous student ID НЕ изменился (UUID v1 тот же — это by design: студент один, когорта меняется).",
        ],
        refs=[
            "⚠ Cohort code limit: 4-16 chars [A-Z0-9-], maxlength=16 в input. Длинные имена (>16) обрезаются при вставке.",
        ])

    add_test_case(doc, "TC-T-10", "Отзыв согласия",
        preconditions="Активная когорта; есть хотя бы 1 upload (TC-T-7).",
        steps=[
            "Открой 📊 → «🗑 Отозвать согласие».",
            "Прочитай предупреждение.",
            "Подтверди отзыв.",
        ],
        expected=[
            "Toast: «Согласие отозвано» (или похожее) + «Локальные данные стёрты».",
            "Модал закрывается.",
            "Повторно открой 📊: Статус = «Выключен».",
            "Anonymous student ID отсутствует в панели (или новый UUID отсутствует пока не дашь consent повторно).",
        ])

    add_test_case(doc, "TC-T-11", "Возврат после отзыва — пустое состояние",
        preconditions="Только что отозвал согласие (TC-T-10).",
        steps=[
            "Закрой 📊 модал.",
            "Перезагрузи страницу (Ctrl+R / Cmd+R).",
            "Открой 📊.",
        ],
        expected=[
            "Статус: «Выключен».",
            "Все опциональные кнопки скрыты, видна только «📊 Дать согласие на участие».",
            "В Транспаренси (если открыть напрямую) — пусто / нет логов.",
        ])

    add_test_case(doc, "TC-T-12", "Re-consent: версия согласия изменилась",
        preconditions=[
            "Активная когорта.",
            "Developer заранее сменил CONSENT_VERSION с 1.0 на 1.1 в public/js/research.js и перезагрузил сервер (см. TC-D-24).",
        ],
        steps=[
            "Перезагрузи страницу с активной когортой.",
            "Открой 📊.",
        ],
        expected=[
            "Статус: «Требуется повторное согласие» (оранжевая отметка).",
            "Видна кнопка «🔄 Обновить согласие».",
            "После нажатия откроется consent screen — версия документа = 1.1.",
            "Anonymous student ID сохранился.",
        ])

    add_test_case(doc, "TC-T-13", "Мобильный UI 375×667",
        preconditions="Открой DevTools → Toggle device toolbar → iPhone SE (375×667). Или открой на реальном мобильном.",
        steps=[
            "Пройди TC-T-1..TC-T-5 заново на мобильной ширине.",
        ],
        expected=[
            "Модал помещается на экране (без горизонтального скролла).",
            "Кнопки нажимаются пальцем (минимум 40px высоты).",
            "Текст читаем без зума.",
        ])

    add_test_case(doc, "TC-T-14", "Планшет / iPad 768×1024",
        preconditions="DevTools → iPad.",
        steps=[
            "Пройди TC-T-1..TC-T-5 на планшете.",
        ],
        expected=[
            "Модал центрирован, не растянут на всю ширину.",
            "Все элементы выровнены, нет artifacts.",
        ])

    add_test_case(doc, "TC-T-15", "Локали RU / EN / HE",
        preconditions="Поддерживается переключение языка в приложении.",
        steps=[
            "Переключи язык на EN, открой 📊, пройди TC-T-2.",
            "Переключи на HE, открой 📊, пройди TC-T-2.",
            "Верни RU.",
        ],
        expected=[
            "EN: все строки переведены, нет '[research.X.Y]'-плейсхолдеров.",
            "HE: текст идёт справа налево (RTL); чекбоксы и кнопки — на корректных позициях.",
            "HE: качество перевода — НЕ оценивается в этом TC (см. Open Q3, deployment blocker).",
            "RU: возврат без артефактов.",
        ],
        refs=[
            "RU/EN ready, HE — machine-grade, нужен native review (docs/ULPAN_RESEARCH_PLAN_v3_2.md §14 Q3).",
        ])

    add_test_case(doc, "TC-T-16", "Offline сценарий: opt-in без сети",
        preconditions="consent ещё не принят.",
        steps=[
            "Выключи Wi-Fi/сеть (или DevTools → Network → Offline).",
            "Открой 📊 → consent → отметь все 5 → «Принимаю».",
            "Попробуй «🔗 Присоединиться к когорте» (любой код).",
            "Включи сеть обратно.",
            "Подожди 5-10 секунд или нажми «⬆ Отправить сейчас».",
        ],
        expected=[
            "Consent сохраняется локально без сети (статус → «Без cohort code» или «Активен»).",
            "Join cohort может пройти ИЛИ показать toast с ошибкой; в любом случае повторный online-attempt должен сработать.",
            "После восстановления сети — upload проходит, в логе появляется запись.",
        ])

    add_test_case(doc, "TC-T-17", "Teacher dashboard: login + пустая когорта",
        preconditions=[
            "Provisioning сделан в Pre-flight §0.3.",
            "Tester имеет доступ к `Smoke-check/test-cohort-credentials.txt` с обоими RESEARCHER_TOKEN_*.",
        ],
        steps=[
            "Открой http://localhost:3000/teacher.html в чистом окне.",
            "В поле Cohort code введи: TEST-PILOT-A.",
            "В поле Researcher token (Bearer) скопируй RESEARCHER_TOKEN_A из credentials.txt.",
            "Нажми «Войти».",
        ],
        expected=[
            "Дашборд открывается без ошибок (нет красного error-баннера).",
            "Header: «TEST-PILOT-A  ·  k=5  ·  schema v1  ·  retain → 2026-06-12» (дата зависит от --retention-days).",
            "Summary tiles: cohort_size = 0, остальные счётчики 0 или «—».",
            "Per-student table → empty state с ⚠ badge «k-anonymity not met (0 < 5)».",
            "Correlations / Scatter → empty state.",
            "Кнопки CSV-экспорта присутствуют (⬇ Aggregates, ⬇ Timeseries, ⬇ Derived).",
        ],
        refs=[
            "Если получил «Ошибка 401 MISSING_BEARER_TOKEN» / 403 BAD_RESEARCHER_TOKEN — проверь что вставил token полностью, без переноса строки. Если 404 — когорта не была provisionirana (вернись к §0.3).",
        ])

    add_test_case(doc, "TC-T-18", "Teacher dashboard: cohort заполнен (k≥5)",
        preconditions=[
            "Pre-flight §0.3 выполнен (npm run smoke:prep). В data/research/ существует SEED-K5 с 12 fake students × 14 days.",
            "Tester имеет доступ к Smoke-check/test-cohort-credentials.txt → RESEARCHER_TOKEN_SEED.",
        ],
        steps=[
            "Открой http://localhost:3000/teacher.html в чистом окне.",
            "Cohort code: SEED-K5.",
            "Researcher token (Bearer): RESEARCHER_TOKEN_SEED из credentials.txt.",
            "Нажми «Войти». Прокрути дашборд сверху вниз.",
        ],
        expected=[
            "Header: «SEED-K5  ·  k=5  ·  schema v1  ·  retain → …».",
            "Cohort overview — 6 tiles с числами > 0: cohort_size ≈ 11-12, days_observed = 14, total_minutes > 0, total_audio > 0, SRS_reviews > 0, notes_created > 0.",
            "Engagement timeline — SVG line chart, 14 точек по дням.",
            "Audio playback / SRS+Notes — два дополнительных графика, видны линии.",
            "Per-student breakdown badge: «k-anonymity met (12 ≥ 5)».",
            "Per-student table: 11-12 строк (один withdrawal case на day 7), сортируется по клику на заголовок.",
            "Outcome correlations: Pearson r показан для каждой метрики vs post_test_score, magnitude label (strong/moderate/weak/none).",
            "Engagement vs exam scatter: dots + trendline.",
        ],
        refs=["Auto-validated by curl на этом этапе: cohort_size=12, k_anonymity_met:true, days_observed=14, students.length=12, per_student_daily.length=12 — см. финальный отчёт."])

    add_test_case(doc, "TC-T-19", "Teacher CSV upload — нормальный сценарий",
        preconditions=[
            "TC-T-18 PASS (dashboard работает).",
            "Pre-flight §0.3 сгенерировал готовый Smoke-check/outcomes-good.csv (12 rows, реальные UUID из SEED-K5).",
        ],
        steps=[
            "В шапке дашборда нажми кнопку «📤 Upload outcomes CSV».",
            "Выбери файл: Smoke-check/outcomes-good.csv.",
            "Подтверди upload.",
            "Дождись response.",
            "Нажми «🔄 Refresh».",
        ],
        expected=[
            "Toast / message: «Outcomes uploaded: { inserted: N, updated: M, total: 12 }».",
            "После refresh: колонка post_test_score заполнена у всех 12 студентов.",
            "Outcome.uploaded_by = «teacher» в каждой строке.",
            "Correlations table заполнен; Scatter показывает 12 точек + trendline.",
        ])

    add_test_case(doc, "TC-T-20", "Teacher CSV upload — sad paths",
        preconditions=[
            "TC-T-18 PASS.",
            "Pre-flight §0.3 создал готовые files: outcomes-bad-no-header.csv + outcomes-bad-empty.csv.",
        ],
        steps=[
            "Шаг 1: Upload Smoke-check/outcomes-bad-no-header.csv (нет header строки).",
            "Шаг 2: Refresh dashboard, проверь что данные не повредились.",
            "Шаг 3: Upload Smoke-check/outcomes-bad-empty.csv (header без data rows).",
            "Шаг 4: Refresh.",
            "(Опционально) Шаг 5: попробуй вообще не выбрать файл и нажать Upload.",
        ],
        expected=[
            "Шаг 1: error toast/banner: «BAD_CSV: CSV header must include 'student_id' (line 1)». Загрузка не сохраняется.",
            "Шаг 2: после refresh — outcomes столбец БЕЗ изменений (значения от TC-T-19 остались).",
            "Шаг 3: error: «NO_ROWS: CSV had a header but no data rows».",
            "Шаг 4: state не изменился.",
            "Шаг 5: либо UI блокирует submit, либо server возвращает 400 EMPTY_BODY.",
            "deletions.log / outcomes.csv на диске НЕ повреждены ни в одном из sad-path сценариев.",
        ],
        refs=[
            "Auto-validated by curl: BAD_CSV / NO_ROWS / EMPTY_BODY / 401 (no token) / 404 (no cohort) — все возвращают 400/401/404, не 500.",
        ])

    add_test_case(doc, "TC-T-21", "Teacher logout + reload",
        preconditions="Активная сессия teacher dashboard.",
        steps=[
            "Нажми ⎋ Logout в шапке.",
            "Перезагрузи страницу.",
        ],
        expected=[
            "Logout → возврат на login screen, поля очищены.",
            "Reload → login screen (не auto-resume на dashboard).",
            "Cohort code и token не сохранились в видимой форме.",
        ])

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────
    # Часть II — Developer
    # ─────────────────────────────────────────────────────────────────────
    add_heading(doc, "Часть II — Developer (DevTools + CLI)",
                level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    add_para(doc,
        "Дополняет Часть I deep-проверками privacy invariants и контрактов API. "
        "Жаргон допустим. Многие шаги — curl + чтение файлов на диске. "
        "На каждом TC обязательно: подтверждение что privacy invariant не нарушен.",
        italic=True, size=10.5, color=RGBColor(0x55, 0x55, 0x55))

    add_test_case(doc, "TC-D-1", "Provision когорты через CLI (повторно после Pre-flight §0.3)",
        preconditions="Pre-flight §0.3 уже выполнил basic provisioning. Этот TC — deep-проверка структуры meta.",
        steps=[
            "Выполни третью когорту для теста: npm run research:cohort -- --code TEST-PILOT-C --retention-days 30 --outcome-scale 0-100 --consent-min 1.0 --k 5",
            "Сохрани plaintext researcher token из вывода.",
            "Просмотри: cat data/research/TEST-PILOT-C/cohort_meta.json",
        ],
        expected=[
            "Файл cohort_meta.json существует.",
            "Поля: code, schema_version=\"v1\", created_at (ISO), k_anonymity_threshold=5, retention_until (ISO), outcome_scale=\"0-100\", researcher_token_hash (sha256, 64 hex chars), consent_version_minimum=\"1.0\".",
            "researcher_token_hash НЕ plaintext (64 chars hex sha256, не base64url).",
            "Plaintext token напечатан ровно один раз в stdout — больше нигде на диске не лежит.",
            "Cleanup: rm -rf data/research/TEST-PILOT-C после TC.",
        ])

    add_test_case(doc, "TC-D-2", "GET aggregates пустой когорты",
        steps=[
            "curl -H \"Authorization: Bearer $RESEARCHER_TOKEN_A\" http://localhost:3000/api/research/v1/cohort/TEST-PILOT-A/aggregates",
        ],
        expected=[
            "HTTP 200.",
            "Top-level JSON: { ok: true, cohort_meta: {...}, cohort_size: 0, k_anonymity_met: false, days_observed: 0, daily_aggregates: [], students: [], per_student_daily: [] }.",
            "cohort_meta содержит: code, schema_version, created_at, k_anonymity_threshold:5, retention_until, outcome_scale:\"0-100\", consent_version_minimum:\"1.0\".",
            "researcher_token_hash в cohort_meta API ответе ОТСУТСТВУЕТ (только на диске).",
        ])

    add_test_case(doc, "TC-D-3", "localStorage до opt-in",
        steps=[
            "Открой клиент в чистом окне.",
            "DevTools → Application → Local Storage → http://localhost:3000.",
        ],
        expected=[
            "Ключей researchEnabled_v1 / researchStudentId_v1 / researchCohortCode_v1 НЕ существует.",
            "researchConsentVersion_v1 НЕ существует.",
        ])

    add_test_case(doc, "TC-D-4", "localStorage после opt-in + join",
        steps=[
            "Пройди consent + join cohort (TC-T-4, TC-T-5).",
            "Application → Local Storage.",
        ],
        expected=[
            "researchEnabled_v1 = \"1\".",
            "researchStudentId_v1 = UUID v4 формат.",
            "researchCohortCode_v1 = ULPAN-DEV-W2026.",
            "researchConsentVersion_v1 = \"1.0\".",
            "researchUploadQueue_v1, researchUploadLog_v1 — JSON массивы (могут быть [] на старте).",
        ])

    add_test_case(doc, "TC-D-5", "Network tab: payload в POST /metrics",
        steps=[
            "DevTools → Network → Fetch/XHR filter.",
            "В клиенте нажми «⬆ Отправить сейчас».",
            "Найди POST /api/research/v1/metrics → Preview / Payload.",
        ],
        expected=[
            "Content-Type: application/json.",
            "Поля payload: schema_format=\"linguistpro-research-v1\", schema_version=\"v1\", student_id, cohort_code, since_ts, upload_ts, app_version, platform, metrics, hours_active.",
            "metrics — ИСКЛЮЧИТЕЛЬНО числовые счётчики и счётные распределения; никаких массивов строк.",
            "В payload ОТСУТСТВУЕТ: any Hebrew/Russian text, note bodies, search query strings, audio blobs, IP, email.",
            "Response: 200 + { stored: true, dedupe: false } (на первом вызове).",
        ])

    add_test_case(doc, "TC-D-6", "Запуск daily aggregator вручную",
        steps=[
            "DevTools → Console.",
            "Выполни: window.LinguistProResearch.runDailyAggregator()",
        ],
        expected=[
            "Promise resolves с объектом результата (stored: bool, dedupe: bool, или skipped).",
            "Если активность за последние сутки = 0 → skipped: true.",
            "Если есть активность → upload проходит, queue не растёт.",
        ])

    add_test_case(doc, "TC-D-7", "Idempotency — повторная отправка",
        steps=[
            "Нажми «⬆ Отправить сейчас» дважды подряд без новых событий.",
            "В Network tab сравни payload first vs second.",
        ],
        expected=[
            "First response: { ok: true, stored: true, dedupe: false, rate_limit_remaining: N }.",
            "Second response: { ok: true, stored: false, dedupe: true, rate_limit_remaining: N-1 } (since_ts/upload_ts совпадают).",
            "На диске data/research/<cohort>/<date>.jsonl — одна строка, не две.",
        ])

    add_test_case(doc, "TC-D-8", "Rate limit — 11+ uploads в день",
        steps=[
            "В Console: for (let i=0; i<12; i++) { await fetch('/api/research/v1/metrics', { method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({...изменённый payload с разным since_ts...}) }); }",
            "Или: симулируй через ручные клики с разным локальным временем.",
        ],
        expected=[
            "Первые 10 запросов: 200.",
            "11-й и далее: 429 RATE_LIMIT с error code.",
            "Клиент-side: после 429 retry-backoff активируется (researchNextRetryAt_v1 заполняется).",
        ])

    add_test_case(doc, "TC-D-9", "Validator: попытка послать raw текст",
        steps=[
            "В Console: window.LinguistProResearch._uploadOnce({ ...валидный payload..., metrics: { ...счётчики..., note_body: 'שלום' } })",
            "Альтернатива через curl: POST с metrics.note_body = «שלום».",
        ],
        expected=[
            "HTTP 400.",
            "Body: { ok: false, error: \"SCHEMA_VIOLATION\", field: \"$.metrics.note_body\", message: \"SCHEMA_VIOLATION: forbidden field \\\"note_body\\\" present\" }.",
            "Сервер не сохраняет такой payload (jsonl не пополнен).",
            "Аналогично: text_content, search_query, audio_bytes, audio_url, username, email, name, phone, ip, geolocation, latitude, longitude, user_agent, device_id, device_serial, timestamp — все banned.",
        ],
        refs=["FORBIDDEN_FIELDS в research/validate.js — полный список."])

    add_test_case(doc, "TC-D-10", "DELETE без cohort_code (full scan, B1 fix)",
        steps=[
            "Сделай self-report outcome (TC-T-8) — это создаст запись в outcomes.csv для твоего UUID.",
            "В Console: const sid = localStorage.getItem('researchStudentId_v1'); fetch(`/api/research/v1/student/${sid}`, { method: 'DELETE' }).then(r => r.json()).then(console.log)",
            "Без `?cohort_code=…` параметра.",
        ],
        expected=[
            "Response 200, body: { ok: true, cohorts_touched: N, records_removed: M } (где N — число cohorts где найден UUID; M — суммарно строк удалено из .jsonl).",
            "Файл data/research/TEST-PILOT-A/outcomes.csv — строка с твоим UUID удалена.",
            "Файл .jsonl — твои строки тоже удалены (wc -l для даты-сегодня = 0).",
            "deletions.log содержит новую append-only запись формата: `<ISO_timestamp> student_id=<uuid> reason=user_withdrawal records_removed=N outcomes_removed=1`.",
        ],
        refs=["B1 fix — privacy commit 32d8cb4. findCohortsForStudent теперь сканирует и outcomes.csv (раньше только .jsonl)."])

    add_test_case(doc, "TC-D-11", "DELETE с несуществующим UUID",
        steps=[
            "curl -X DELETE \"http://localhost:3000/api/research/v1/student/00000000-0000-4000-8000-000000000000?cohort_code=TEST-PILOT-A\"",
        ],
        expected=[
            "Response 200 (idempotent), body: { ok: true, cohorts_touched: 1 (или 0), records_removed: 0 }.",
            "На диске никаких изменений.",
            "deletions.log — может пополниться записью с records_removed=0 (audit best practice).",
            "В НИКАКОМ случае не 500-ка.",
        ])

    add_test_case(doc, "TC-D-12", "Privacy invariant — нет hebrew chars в research-data/",
        steps=[
            "После прогона активного использования (несколько uploads):",
            "grep -P '[\\x{0590}-\\x{05FF}]' data/research/ -r 2>/dev/null  (или эквивалент на Windows через Select-String)",
            "Также: grep -i 'note_body\\|search_query\\|content' data/research/ -r",
        ],
        expected=[
            "Нет совпадений с hebrew unicode.",
            "Нет ключей note_body / search_query / content в JSON payload'ах.",
            "outcomes.csv — только числа и UUID, никакого PII.",
        ])

    add_test_case(doc, "TC-D-13", "Что лежит на сервере — manual JSONL inspection",
        steps=[
            "cat data/research/ULPAN-DEV-W2026/<today>.jsonl",
            "Распарси через jq: jq '.' data/research/ULPAN-DEV-W2026/<today>.jsonl",
        ],
        expected=[
            "Каждая строка — валидный JSON.",
            "Поля строго соответствуют RESEARCH_METRICS_SCHEMA.md §10.",
            "metrics — только счётчики; нет вложенных строковых массивов.",
        ])

    add_test_case(doc, "TC-D-14", "k-anonymity gate — когорта с 4 студентами",
        steps=[
            "Создай когорту: npm run research:cohort -- --code K-TEST-4 --k 5",
            "Через curl или client отправь 4 different student uploads (UUID v4 каждый).",
            "GET aggregates с токеном.",
        ],
        expected=[
            "Response 200.",
            "cohort_size: 4.",
            "k_anonymity_met: false.",
            "students: [] (length 0).",
            "per_student_daily: [] (length 0).",
            "daily_aggregates — заполнены cohort-wide агрегатами.",
            "Per-student detail скрыт.",
        ])

    add_test_case(doc, "TC-D-15", "k-anonymity gate — когорта с 5 студентами",
        steps=[
            "В той же когорте K-TEST-4 отправь 5-й upload от нового UUID.",
            "GET aggregates.",
        ],
        expected=[
            "Response 200.",
            "cohort_size: 5.",
            "k_anonymity_met: true.",
            "students: array length 5 — per-student breakdowns теперь видны.",
            "per_student_daily: array length 5.",
        ])

    add_test_case(doc, "TC-D-16", "GET aggregates без токена",
        steps=[
            "curl http://localhost:3000/api/research/v1/cohort/TEST-PILOT-A/aggregates",
        ],
        expected=[
            "HTTP 401.",
            "Тело: { ok: false, error: \"MISSING_BEARER_TOKEN\" }.",
        ])

    add_test_case(doc, "TC-D-17", "GET aggregates с неверным токеном (и cross-cohort token)",
        steps=[
            "curl -H \"Authorization: Bearer wrong-token-here\" http://localhost:3000/api/research/v1/cohort/TEST-PILOT-A/aggregates",
            "Затем: curl -H \"Authorization: Bearer $TOKEN_B\" http://localhost:3000/api/research/v1/cohort/TEST-PILOT-A/aggregates",
        ],
        expected=[
            "Wrong token: HTTP 403, body: { ok: false, error: \"BAD_RESEARCHER_TOKEN\" }.",
            "Cross-cohort (token от B применён к A): HTTP 403 BAD_RESEARCHER_TOKEN — token строго привязан к когорте.",
            "Сервер не возвращает hint о валидности.",
        ])

    add_test_case(doc, "TC-D-18", "Token rotation Procedure B (in-place)",
        preconditions="Когорта существует, текущий token известен.",
        steps=[
            "Сгенерируй новый plaintext: NEW_TOKEN=$(node -e \"console.log(require('crypto').randomBytes(24).toString('base64url'))\")",
            "Compute sha256: NEW_HASH=$(node -e \"console.log(require('crypto').createHash('sha256').update('$NEW_TOKEN').digest('hex'))\")",
            "Перепиши cohort_meta.json researcher_token_hash = $NEW_HASH.",
            "Verify: GET с старым токеном → 401; GET с новым токеном → 200.",
        ],
        expected=[
            "Старый token немедленно перестаёт работать.",
            "Новый token работает.",
            "cohort_meta.json valid JSON после правки.",
        ],
        refs=["RESEARCHER_GUIDE.md §2.1.1 Procedure B"])

    add_test_case(doc, "TC-D-19", "CSV outcomes — корректный upload",
        steps=[
            "В teacher.html upload outcomes.csv с заголовком student_id,pre_test_score,post_test_score,exam_date,uploaded_by + 3 строки.",
            "Проверь data/research/<cohort>/outcomes.csv.",
        ],
        expected=[
            "Файл существует.",
            "Header валиден.",
            "Строки совпадают с upload (один-к-одному).",
            "При refresh dashboard — post_test_score появляется в per-student table.",
        ])

    add_test_case(doc, "TC-D-20", "Outcome withdrawal cleanup (B1 регрессия)",
        steps=[
            "Создай UUID-A в когорте, отправь self-report outcome (TC-T-8).",
            "Проверь outcomes.csv содержит UUID-A.",
            "Из клиента того же UUID-A — withdrawal (TC-T-10).",
            "Снова проверь outcomes.csv.",
        ],
        expected=[
            "Строка с UUID-A удалена из outcomes.csv после withdrawal.",
            "deletions.log содержит запись с outcomes_removed=1.",
            "consent template promise («удаление всех ваших ранее загруженных данных») выполнен.",
        ],
        refs=["B1 fix — commit 32d8cb4"])

    add_test_case(doc, "TC-D-21", "Audit log deletions.log",
        steps=[
            "После TC-D-10 и TC-D-20: cat data/research/<cohort>/deletions.log",
        ],
        expected=[
            "Каждая запись: <ISO timestamp> student_id=<uuid> reason=user_withdrawal records_removed=N outcomes_removed=M",
            "Никакого PII (имена, IP, etc).",
            "Append-only (никаких rewrite).",
        ])

    add_test_case(doc, "TC-D-22", "Retry backoff при недоступном сервере",
        steps=[
            "Активируй research mode + join cohort.",
            "Останови сервер (Ctrl+C).",
            "В Console: window.LinguistProResearch.runDailyAggregator()",
            "Проверь researchUploadQueue_v1 и researchNextRetryAt_v1 в Application.",
            "Запусти сервер обратно.",
            "Жди / нажми «⬆ Отправить сейчас» вручную.",
        ],
        expected=[
            "При недоступном сервере: queue растёт (max 30), nextRetryAt установлен в future.",
            "После восстановления: queue flush'ится, server stores все pending uploads.",
            "Idempotency dedupes повторные попытки.",
        ])

    add_test_case(doc, "TC-D-23", "Multi-device — два UUID в одной когорте",
        steps=[
            "Окно 1: пройди TC-T-4 + TC-T-5 (UUID-A).",
            "Окно 2 (приватное / другой профиль): пройди TC-T-4 + TC-T-5 с тем же cohort code (UUID-B).",
            "GET aggregates.",
        ],
        expected=[
            "В per-student table: 2 разных UUID.",
            "Это by design — manual linking deferred к v3.3.",
        ],
        refs=["docs/ULPAN_RESEARCH_PLAN_v3_2.md §14 Q6"])

    add_test_case(doc, "TC-D-24", "CONSENT_VERSION bump → re-consent",
        steps=[
            "Останови сервер.",
            "В public/js/research.js: const CONSENT_VERSION = '1.0' → '1.1'.",
            "Запусти сервер обратно, hard refresh клиента.",
            "Открой 📊 (TC-T-12).",
        ],
        expected=[
            "Статус: «Требуется повторное согласие».",
            "Anonymous student ID не сменился.",
            "После повторного consent — researchConsentVersion_v1 = «1.1».",
            "После теста: верни CONSENT_VERSION = «1.0» (если правил локально), иначе изменение нужно докоммитить как v3.3 epic.",
        ],
        refs=["Q2 — ad-hoc rule, docs/ULPAN_RESEARCH_PLAN_v3_2.md §14"])

    add_test_case(doc, "TC-D-25", "Auto-smoke runner",
        steps=[
            "npm run smoke:research:fast  (только функциональные suites, ~4s)",
            "npm run smoke:research        (полный с screenshot'ами, ~8s)",
            "Просмотри 9 PNG в Smoke-check/teacher-dashboard/<timestamp>/.",
        ],
        expected=[
            "Server smoke: 25/25.",
            "Browser smoke: 21/21.",
            "Teacher smoke: 14/14.",
            "Screenshots: 9 PNG, нет «empty state» там, где должны быть числа.",
            "Сравни новые PNG с baseline (commit e8df486) — visual diff минимален.",
        ],
        refs=[
            "⚠ Известный flake: первый прогон после простоя иногда фейлит teacher-smoke с Timeout на `.summary-tile`. Перепрогон обычно green. Если 2+ прогона подряд fail — это уже регрессия, фиксируй.",
        ])

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────
    # Sign-off
    # ─────────────────────────────────────────────────────────────────────
    add_heading(doc, "Sign-off и закрытие smoke",
                level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    add_para(doc,
        "Прогон считается полным когда:")
    add_bullets(doc, [
        "Все 21 Tester TC отмечены PASS или N/A с объяснением.",
        "Все 25 Developer TC отмечены PASS или N/A с объяснением.",
        "npm run smoke:research возвращает 60/60 + 9 PNG green.",
        "Privacy invariants (TC-D-12) — без нарушений.",
        "B1 регрессия (TC-D-20) — PASS.",
        "HE consent native review — отдельно отмечен как deployment blocker, не блокирует pilot.",
    ])

    add_para(doc, "")
    t = doc.add_table(rows=4, cols=2)
    rows = [
        ("Тестер (имя)", "_______________________________"),
        ("Дата прогона", "_______________________________"),
        ("Среда (local / staging / preview-deploy)", "_______________________________"),
        ("Подпись / approval", "_______________________________"),
    ]
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v
        set_cell_shading(t.cell(i, 0), 'EAF1F8')
        for p in t.cell(i, 0).paragraphs:
            for rr in p.runs:
                rr.bold = True

    add_para(doc, "")
    add_para(doc,
        "После полностью green smoke — разрешение на distribution когорты "
        "2-3 friendly users (см. docs/PARALLEL_WORK_PLAN_DURING_PILOT.md).",
        italic=True, color=RGBColor(0x88, 0x44, 0x00))

    add_heading(doc, "Cleanup после smoke", level=2)
    add_bullets(doc, [
        "Одной командой: `npm run smoke:prep:clean` — удалит cohorts + outcomes-*.csv + credentials.txt.",
        "Если правил CONSENT_VERSION в research.js (TC-T-12 / TC-D-24) — `git checkout public/js/research.js`.",
        "Останови dev-сервер: Ctrl+C.",
    ])

    add_heading(doc, "Revision log", level=2)
    add_bullets(doc, [
        "r1 — 2026-05-13: initial release вместе с v3.2.0 mega-release.",
        "r2 — 2026-05-14: исправлено 3 блокера из первого пилотного прогона. "
        "(B1) TC-T-6 переписан под реальный uploads-only UI. "
        "(B2) Тестовые cohort codes уложены в maxlength=16 (TEST-PILOT-A / -B). "
        "(B3) Pre-flight §0.3 добавлен — manual provisioning + credentials.txt. "
        "Также точные response форматы TC-D-2/7/9/10/14/15/16/17.",
        "r3 — 2026-05-14: zero-manual-prep после второго пилотного прогона. "
        "(P1) `npm run smoke:prep` готовит ВСЁ одной командой: TEST-PILOT-A/B + seeded SEED-K5 + credentials.txt + outcomes-good.csv + outcomes-bad-*.csv. "
        "(P2) TC-T-18/19/20 переписаны под готовые артефакты — тестер просто указывает путь к файлу. "
        "(P3) `npm run smoke:prep:clean` для очистки. "
        "(P4) Прогон Developer-секции автоматизирован через curl автором; TC-T-18, 19, 20 + все TC-D-* кроме UI-only — auto-PASSED.",
    ])

    # ── References ───────────────────────────────────────────────────────
    add_page_break(doc)
    add_heading(doc, "Приложение — Где живут источники истины",
                level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    refs_table = doc.add_table(rows=1, cols=2)
    refs_table.style = 'Light Grid Accent 1'
    refs_table.rows[0].cells[0].text = "Что"
    refs_table.rows[0].cells[1].text = "Где"
    refs = [
        ("Полная privacy policy", "docs/PRIVACY.md"),
        ("Wire contract / schema", "docs/RESEARCH_METRICS_SCHEMA.md"),
        ("Consent templates (RU / EN / HE skeleton)", "docs/RESEARCH_ETHICS_CONSENT_TEMPLATE.md"),
        ("Researcher operational guide", "docs/RESEARCHER_GUIDE.md"),
        ("Master plan + open questions", "docs/ULPAN_RESEARCH_PLAN_v3_2.md"),
        ("Client logic (research mode)", "public/js/research.js"),
        ("Client UX", "public/js/research-ui.js"),
        ("Teacher dashboard", "public/teacher.html + public/js/teacher.js"),
        ("Server endpoints", "research/storage.js + research/validate.js + server routes"),
        ("Auto-smoke runners", "scripts/research/all-smoke.js (+ individual sub-runners)"),
        ("Baseline visual regression", "Smoke-check/teacher-dashboard/ (commit e8df486)"),
        ("Release tag", "v3.2.0 (commit 32d8cb4)"),
    ]
    for k, v in refs:
        r = refs_table.add_row().cells
        r[0].text = k
        r[1].text = v

    # Save
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_FILE))
    print(f"OK: {OUT_FILE}")


if __name__ == "__main__":
    build_doc()
